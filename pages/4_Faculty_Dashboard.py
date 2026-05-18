import streamlit
import streamlit as st
from components.style import load_css
import os
import json
import io
from typing import List
from pydantic import BaseModel, Field
from crewai import Agent, Task, Crew, LLM
from langchain_community.document_loaders import PyPDFLoader
from dotenv import load_dotenv
from docx import Document

load_css()

with st.sidebar:
    st.image("assets/logo.png", width=150)
    st.markdown("## 👨‍🏫 Faculty Panel")
    menu = st.radio("Navigation", [
        "Dashboard",
        "Generate MCQ",
        "Generate LAQ",
        "Create Test",
        "Logout"
    ])

st.title("Faculty Dashboard")

if menu == "Generate MCQ":
    # --- 1. Structured Data Schemas ---
    class MCQQuestion(BaseModel):
        question: str = Field(..., description="The multiple choice question text.")
        options: List[str] = Field(...,
                                   description="Exactly 4 options, formatted like ['A) Option 1', 'B) Option 2', 'C) Option 3', 'D) Option 4'].")
        correct_option: str = Field(..., description="The correct option letter index (e.g., 'A', 'B', 'C', or 'D').")
        explanation: str = Field(..., description="A brief, high-quality explanation of why this option is correct.")


    class MCQExam(BaseModel):
        questions: List[MCQQuestion]


    load_dotenv()

    # Initialize Gemini Model
    gemini_flash = LLM(
        model="gemini/gemini-3.1-flash-lite-preview",
        temperature=0.4,
        api_key=os.getenv("GOOGLE_API_KEY")
    )


    # --- 2. DOCX Generation Utilities ---
    def generate_questions_docx(exam_data):
        doc = Document()
        doc.add_heading('MCQ Assessment', level=1)
        doc.add_paragraph("Answer all the questions below by selecting the most appropriate option.\n")

        for i, q in enumerate(exam_data['questions']):
            doc.add_paragraph(f"{q['question']}", style='List Number')
            for option in q['options']:
                doc.add_paragraph(f"    {option}")
            doc.add_paragraph("")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


    def generate_answers_docx(exam_data):
        doc = Document()
        doc.add_heading('Answer Key', level=1)
        doc.add_paragraph("Review the correct options along with contextual logic below.\n")

        for i, q in enumerate(exam_data['questions']):
            p = doc.add_paragraph()
            p.add_run(f"Q{i + 1}. Correct Answer: ").bold = True
            p.add_run(f"[{q['correct_option']}]\n")
            p.add_run("Explanation: ").italic = True
            p.add_run(f"{q['explanation']}\n")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


    # --- 3. Streamlit UI Configuration ---
    st.set_page_config(page_title="AI MCQ Engine", layout="centered")
    st.title("🎯 AI Multiple Choice Generator")
    st.write("Upload up to 4 PDFs and customize the question count and difficulty for each individually.")

    # Persistent storage for generation output
    if "generated_mcqs" not in st.session_state:
        st.session_state.generated_mcqs = None

    # Sidebar Controls
    with st.sidebar:
        st.header("⚙️ Global Actions")
        if st.button("🔄 Reset Generator"):
            # Explicit reset of dynamic file parameters
            if "mcq_pdf_uploader" in st.session_state and st.session_state.mcq_pdf_uploader:
                for file in st.session_state.mcq_pdf_uploader:
                    if f"num_qs_{file.name}" in st.session_state:
                        st.session_state[f"num_qs_{file.name}"] = 5
                    if f"diff_{file.name}" in st.session_state:
                        st.session_state[f"diff_{file.name}"] = 5
            st.session_state.generated_mcqs = None
            st.rerun()

    # Document Upload Section (Max 4 PDFs)
    uploaded_files = st.file_uploader(
        "Upload Source Text (Maximum 10 PDFs)",
        type=["pdf"],
        accept_multiple_files=True,
        key="mcq_pdf_uploader"
    ) or []

    # Configuration Tracking List
    pdf_configs = []

    if uploaded_files:
        if len(uploaded_files) > 10:
            st.error("⚠️ Maximum capacity exceeded. Please upload up to 10 PDFs only.")
        else:
            st.write("### 🛠️ Per-PDF Customization")

            # Dynamically build control inputs for each uploaded document
            for idx, file in enumerate(uploaded_files):
                with st.expander(f"📄 Document {idx + 1}: {file.name}", expanded=True):
                    col1, col2 = st.columns(2)
                    with col1:
                        n_qs = st.number_input(
                            "Number of MCQs",
                            min_value=1,
                            max_value=30,
                            value=5,
                            step=1,
                            key=f"num_qs_{file.name}"
                        )
                    with col2:
                        diff = st.slider(
                            "Difficulty Profile",
                            1, 10, 5,
                            key=f"diff_{file.name}",
                            help="1 = Fundamental Recall, 10 = Deep Conceptual Analysis"
                        )

                    # Cache the individual requirements mapped to the file object
                    pdf_configs.append({
                        "file_obj": file,
                        "num_questions": n_qs,
                        "difficulty": diff,
                        "temp_path": f"temp_mcq_source_{idx}.pdf"
                    })

    # --- 4. Processing & Core AI Execution Loop ---
    if pdf_configs:
        if st.button("🚀 Synthesize Customized MCQ Exam"):
            master_questions_list = []

            with st.status("Processing Documents Separately...", expanded=True) as status:

                # Loop through each PDF and its custom configuration profile
                for config in pdf_configs:
                    file = config["file_obj"]
                    status.write(
                        f"📖 Extracting & generating {config['num_questions']} Qs (Difficulty: {config['difficulty']}/10) from `{file.name}`...")

                    # Write temp file for the loader
                    with open(config["temp_path"], "wb") as f:
                        f.write(file.getbuffer())

                    loader = PyPDFLoader(config["temp_path"])
                    pages = loader.load()
                    file_content = "\n".join([page.page_content for page in pages])

                    # Define Agents
                    analyst = Agent(
                        role='Curriculum Analyst',
                        goal='Distill foundational concepts and testable items from raw text.',
                        backstory='Specialized in academic taxonomy mapping.',
                        llm=gemini_flash
                    )
                    psychometrist = Agent(
                        role='Psychometric Assessment Officer',
                        goal=f"Draft exactly {config['num_questions']} clear multiple choice questions.",
                        backstory='Expert in constructing rigorous diagnostic academic assessments.',
                        llm=gemini_flash
                    )

                    extraction_task = Task(
                        description=f"Analyze core themes within this document text: {file_content}",
                        expected_output="An outline of key testable objectives.",
                        agent=analyst
                    )

                    generation_task = Task(
                        description=f"""
                        Using the text and analysis objectives, generate exactly {config['num_questions']} multiple choice questions.

                        SPECIFIC FILE RULES:
                        - Target Difficulty Profile: {config['difficulty']}/10
                        - Every question must include exactly 4 choices labeled A), B), C), and D).
                        - Distractors must be plausible but definitively incorrect based strictly on the source text.
                        - Provide a concise clarification for the correct solution.
                        """,
                        expected_output="JSON MCQExam object.",
                        agent=psychometrist,
                        context=[extraction_task],
                        output_json=MCQExam
                    )

                    # Process the individual file configuration job
                    file_crew = Crew(agents=[analyst, psychometrist], tasks=[extraction_task, generation_task])
                    file_output = file_crew.kickoff()

                    # Append individual results to the master collection
                    file_data = file_output.to_dict()
                    master_questions_list.extend(file_data.get('questions', []))

                    # Clean up local file system path tracking
                    if os.path.exists(config["temp_path"]):
                        os.remove(config["temp_path"])

                # Commit total aggregated exam paper back to session state
                st.session_state.generated_mcqs = {"questions": master_questions_list}
                status.update(label="✅ Comprehensive Customized Exam Ready!", state="complete")

    # --- 5. Clean Rendering Block ---
    if st.session_state.generated_mcqs:
        st.divider()
        st.subheader("📋 Plain Text Document Preview")
        st.caption(f"Total Questions Generated: {len(st.session_state.generated_mcqs['questions'])}")

        # Render Plain Text Screen Feed Loop
        for idx, item in enumerate(st.session_state.generated_mcqs['questions']):
            st.markdown(f"**Q{idx + 1}. {item['question']}**")
            for option in item['options']:
                st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;{option}")
            st.write("")

        st.divider()

        with st.sidebar:
            question_paper_buffer = generate_questions_docx(st.session_state.generated_mcqs)
            st.download_button(
                label="📥 Download Custom Questions (.DOCX)",
                data=question_paper_buffer,
                file_name="custom_mcq_questions.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            answer_key_buffer = generate_answers_docx(st.session_state.generated_mcqs)
            st.download_button(
                label="📥 Download Custom Answer Key (.DOCX)",
                data=answer_key_buffer,
                file_name="custom_mcq_answers.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

elif menu == "Generate LAQ" :
    # --- 1. Structured Data Schemas ---
    class LAQQuestion(BaseModel):
        question: str = Field(..., description="The descriptive, open-ended long answer question text.")
        mark_value: int = Field(..., description="The marks assigned to this specific question.")
        key_points: List[str] = Field(...,
                                      description="Essential conceptual milestones, core theories, or facts that must be present in a flawless response.")
        grading_rubric: str = Field(...,
                                    description="A clear, structural guide detailing how an evaluator should award points based on the mark value.")


    class LAQExam(BaseModel):
        questions: List[LAQQuestion]


    load_dotenv()

    # Initialize Gemini Model
    gemini_flash = LLM(
        model="gemini/gemini-3.1-flash-lite-preview",
        temperature=0.5,
        api_key=os.getenv("GOOGLE_API_KEY")
    )


    # --- 2. DOCX Generation Utilities ---
    def generate_questions_docx(exam_data):
        doc = Document()
        doc.add_heading('Descriptive Examination Paper', level=1)
        doc.add_paragraph("Instructions: Read each item carefully and provide exhaustive, structured answers.\n")

        for i, q in enumerate(exam_data['questions']):
            p = doc.add_paragraph()
            p.add_run(f"Q{i + 1}. ").bold = True
            p.add_run(f"{q['question']} ")
            p.add_run(f"[{q['mark_value']} Marks]").bold = True
            doc.add_paragraph("")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


    def generate_rubric_docx(exam_data):
        doc = Document()
        doc.add_heading('Evaluation Rubric & Key Guidelines', level=1)
        doc.add_paragraph("This guide establishes standard benchmarks for assessing student submissions.\n")

        for i, q in enumerate(exam_data['questions']):
            doc.add_heading(f"Question {i + 1} Rubric ({q['mark_value']} Marks)", level=2)

            p_q = doc.add_paragraph()
            p_q.add_run("Target Prompt: ").italic = True
            p_q.add_run(q['question'])

            doc.add_heading("Expected Key Points:", level=3)
            for point in q['key_points']:
                doc.add_paragraph(f"• {point}")

            doc.add_heading("Grading Breakdown Criteria:", level=3)
            doc.add_paragraph(q['grading_rubric'])
            doc.add_paragraph("—" * 20)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


    # --- 3. Streamlit UI Configuration ---
    st.set_page_config(page_title="AI LAQ Engine", layout="centered")
    st.title("📑 AI LAQ Generator")
    st.write("Construct descriptive assessments from up to 10 source documents with file-specific blueprints.")

    # Persistent storage for generation output
    if "generated_laqs" not in st.session_state:
        st.session_state.generated_laqs = None

    # Document Upload Section (Max 10 PDFs)
    uploaded_files = st.file_uploader(
        "Upload Source Materials (Maximum 10 PDFs)",
        type=["pdf"],
        accept_multiple_files=True,
        key="laq_pdf_uploader"
    ) or []

    # Configuration Tracking List
    pdf_configs = []

    if uploaded_files:
        if len(uploaded_files) > 10:
            st.error("⚠️ Maximum capacity exceeded. Please limit your selection to 10 PDFs or fewer.")
        else:
            st.write("### 🛠️ Per-PDF Customization")

            # Dynamically build specific requirement modules for each uploaded asset
            for idx, file in enumerate(uploaded_files):
                with st.expander(f"📄 Document {idx + 1}: {file.name}", expanded=True):

                    diff = st.slider(
                        "Difficulty Profile",
                        1, 10, 5,
                        key=f"diff_{file.name}",
                        help="1 = Fundamental Explanations, 10 = Deep Critical Case Syntheses"
                    )

                    # Dynamic Blueprint Map (1-10 marks expansion grid)
                    mark_distribution = {}
                    with st.expander("📊 Assign Question Counts (1-10 Marks)", expanded=False):
                        col1, col2 = st.columns(2)
                        for m in range(1, 11):
                            target_col = col1 if m <= 5 else col2
                            with target_col:
                                count = st.number_input(
                                    f"{m} Mark Qs:",
                                    min_value=0,
                                    max_value=20,
                                    value=0,
                                    step=1,
                                    key=f"marks_{m}_{file.name}"
                                )
                                if count > 0:
                                    mark_distribution[m] = count

                    # --- NEW FEAT: Per-PDF Statistics Breakdown ---
                    total_file_qs = sum(mark_distribution.values())
                    total_file_marks = sum(m * count for m, count in mark_distribution.items())

                    if total_file_qs > 0:
                        st.markdown("---")
                        stat_col1, stat_col2 = st.columns(2)
                        with stat_col1:
                            st.metric(label="Selected Questions", value=total_file_qs)
                        with stat_col2:
                            st.metric(label="Total Marks Value", value=f"{total_file_marks} M")

                    # Cache parameters only if questions are requested for this file
                    if total_file_qs > 0:
                        pdf_configs.append({
                            "file_obj": file,
                            "mark_distribution": mark_distribution,
                            "total_questions": total_file_qs,
                            "difficulty": diff,
                            "temp_path": f"temp_laq_source_{idx}.pdf"
                        })

    # --- 4. Processing & Core AI Execution Loop ---
    if pdf_configs:
        st.write("")
        if st.button("🚀 Synthesize Customized LAQ Exam"):
            master_questions_list = []

            with st.status("Processing Documents Separately...", expanded=True) as status:

                for config in pdf_configs:
                    file = config["file_obj"]
                    status.write(f"📖 Generating {config['total_questions']} LAQs from `{file.name}`...")

                    with open(config["temp_path"], "wb") as f:
                        f.write(file.getbuffer())

                    loader = PyPDFLoader(config["temp_path"])
                    pages = loader.load()
                    file_content = "\n".join([page.page_content for page in pages])

                    # --- NEW AGENT TEAM: Splitting Question Authoring vs Rubric Design ---
                    analyst = Agent(
                        role='Structural Academic Researcher',
                        goal='Deconstruct complex arguments, contextual nuances, and analytical paths from reference materials.',
                        backstory='Specialized in itemizing core curriculum specifications for senior assessments.',
                        llm=gemini_flash
                    )

                    examiner = Agent(
                        role='Chief Evaluation Assessor',
                        goal="Draft rigorous, clear, open-ended descriptive questions and identify core factual key points.",
                        backstory='An expert academic writer skilled at formulating balanced assessment prompts tailored to specific mark criteria.',
                        llm=gemini_flash
                    )

                    rubric_specialist = Agent(
                        role='Educational Psychometrist & Rubric Architect',
                        goal='Design precise grading rubrics, multi-tiered point allocation maps, and strict evaluation metrics.',
                        backstory='A veteran academic standardizations evaluator who knows how to penalize fluff/keyword stuffing and reward genuine synthesis.',
                        llm=gemini_flash
                    )

                    extraction_task = Task(
                        description=f"Isolate broad conceptual mechanisms and multi-layered topics within this document text: {file_content}",
                        expected_output="An analytical summary matching potential long-form assessment pathways.",
                        agent=analyst
                    )

                    # Focus completely on question prompt formulation and core key facts
                    generation_task = Task(
                        description=f"""
                        Using the text analysis overview, generate descriptive questions matching this exact blueprint map: {json.dumps(config['mark_distribution'])}.
                        Total questions to generate from this document: {config['total_questions']}.

                        STRICT EXECUTION PROTOCOLS:
                        - Target Difficulty Profile: {config['difficulty']}/10
                        - You must generate exactly the number of questions specified for each mark tier in the blueprint map.
                        - Formulate precise, open-ended questions and list out the foundational 'key_points' required for each answer.
                        - Leave the 'grading_rubric' property completely empty ("") for the next agent to fill.
                        """,
                        expected_output="JSON structure containing draft questions and core key points.",
                        agent=examiner,
                        context=[extraction_task]
                    )

                    # Delegate evaluation mapping entirely to the specialized rubric agent
                    rubric_task = Task(
                        description=f"""
                        Review the draft questions and matching key points generated in the previous step.
                        For each question, draft an exhaustive 'grading_rubric' demonstrating exactly how an evaluator should distribute the points fairly up to its assigned max mark value.

                        RUBRIC REQUIREMENTS:
                        1. Clearly establish what constitutes full marks vs partial credit.
                        2. Address structural, logic flow, and clarity requirements proportional to the marks assigned.
                        """,
                        expected_output="JSON LAQExam object with fully completed rubrics.",
                        agent=rubric_specialist,
                        context=[generation_task],
                        output_json=LAQExam
                    )

                    # Run the multi-stage pipeline for this file
                    file_crew = Crew(
                        agents=[analyst, examiner, rubric_specialist],
                        tasks=[extraction_task, generation_task, rubric_task]
                    )
                    file_output = file_crew.kickoff()

                    file_data = file_output.to_dict()
                    master_questions_list.extend(file_data.get('questions', []))

                    if os.path.exists(config["temp_path"]):
                        os.remove(config["temp_path"])

                st.session_state.generated_laqs = {"questions": master_questions_list}
                status.update(label="✅ Comprehensive Evaluation Paper Compiled!", state="complete")

    # --- 5. Clean Rendering Block ---
    laq_data = st.session_state.generated_laqs

    if isinstance(laq_data, dict) and "questions" in laq_data:
        st.divider()
        st.subheader("📋 Document Preview")
        st.caption(f"Total Descriptive Questions Generated: {len(laq_data['questions'])}")

        for idx, item in enumerate(laq_data['questions']):
            st.markdown(f"**Q{idx + 1}. {item['question']}**  \n*[{item['mark_value']} Marks]*")
            st.write("")

        # --- 6. Sidebar Panel (Actions & Downloads) ---
    with st.sidebar:
        st.header("⚙️ System Control")
        if st.button("🔄 New LAQ", use_container_width=True):
            st.session_state.clear()
            st.rerun()

        # Context-Aware Export Section
        if isinstance(laq_data, dict) and "questions" in laq_data:
            st.write("")
            st.divider()
            st.header("📥 Export Panel")
            st.info(f"📁 {len(laq_data['questions'])} questions ready for download.")

            question_paper_buffer = generate_questions_docx(laq_data)
            st.download_button(
                label="📄 Download Questions (.DOCX)",
                data=question_paper_buffer,
                file_name="laq_question_paper.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

            rubric_buffer = generate_rubric_docx(laq_data)
            st.download_button(
                label="🔑 Download Evaluation Rubric (.DOCX)",
                data=rubric_buffer,
                file_name="laq_assessment_rubric.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
# 🚫 Block direct access
_='''if "authenticated" not in st.session_state or not st.session_state["authenticated"]:
    st.warning("Please login first.")
    st.switch_page("pages/1_Login.py")

# 🚫 Role protection
if st.session_state.get("role") != "Faculty":
    st.error("Access Denied.")
    st.stop()
    
if st.sidebar.button("Logout"):
    st.session_state.clear()
    st.switch_page("app.py")'''